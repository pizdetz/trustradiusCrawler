# Author: Colton Sweeney
# Date: August 10, 2017
# Name: TrustRadius Review Scraper
#
# Lang: Python
# Dependencies: BeautifulSoup4
#               Requests
#               python-docx
#
# Description: As part of my current workload, I have to create
#              weekly reports gathered from reviews on TrustRadius.
#              In order to expedite this process, I decided to create 
#              simple Python bot that scours the TrustRadius page for
#              all new reviews and then creates detailed reports directly 
#              from the information gathered in those reports.

# import required dependencies
from bs4 import BeautifulSoup
import requests
import docx
from datetime import date
from operator import itemgetter, attrgetter, methodcaller
from Review import Review
import sys

reviewNum = input("How many reviews would you like?\nEnter a number from 0 to 25: \n")

# Open up TrustRadius page for Sitefinity
url = "https://www.trustradius.com/products/progress-sitefinity/reviews"
data = requests.get(url).text
soup = BeautifulSoup(data, "html.parser")

# Find all review links from main product page
query = soup.find_all('a', {
    'class' : 'link-to-review-btn'
})

# List that holds URLs of individual reviews
links = []

# Append all the URLs to the list we created
for link in query:
    links.append('https://www.trustradius.com' + link.get('href'))

# Print number of reviews to command line (if used for plural/singular case)
if reviewNum == 1:
    print "%d links found. Creating %d review." % (len(links), reviewNum)
else:
    print "%d links found. Creating %d reviews." % (len(links), reviewNum)


# Function for returning review sections from review page
# return (dictionary): a key-value list of the headings and review text
# parameters: link (string): the url for the review
def findMaterials(link):
    # Parse the given link into some Beautiful Soup
    req = requests.get(link).text
    reviews = BeautifulSoup(req, 'html.parser')

    # Set up list string variables.
    reviewAuthor = []
    reviewPosition = []
    reviewCompany = []
    reviewRating = []
    sectionHeading = []
    sectionText = []
    sectionDate = ''


    # Find the authors name (if there is one)
    for review in reviews.find_all('span', {'itemprop': 'author'}):
        reviewAuthor.append(review.contents[0].text)
    
    # Find the author's position and company (if applicable)
    for review in reviews.find_all('span', {'class': 'user-info'}):
        reviewPosition.append(review.contents[0].text)
        reviewCompany.append(review.contents[1].text)

    # Find what the user rated Sitefinity
    reviewRating = reviews.find_all('span', class_='number')[0].text

    # Perform find.contents[] for all of the headings and text
    # and append them to our functions variables
    for review in reviews.find_all('div', {'class': 'description'}):
        
        # Receive review section headings
        for head in range(6):
            sectionHeading.append(review.contents[0].contents[0].contents[1].contents[head].contents[0].contents[0].contents[0].contents[0].text)

        # Receive review section bodies
        for body in range(6):
            sectionText.append(" %s" % review.contents[0].contents[0].contents[1].contents[body].contents[1].contents[0].contents[0].contents[0].text)


    # Wrap up the review information into a dictionary, this is for easy handling    
    reviewDict = dict(zip(sectionHeading, sectionText))

    # Get's the date of the review from the review's URL
    sectionDate = link[56:-9]
    days = date(int(sectionDate[:-6]), int(sectionDate[5:-3]), int(sectionDate[8:]))

    # Create a new review using our Review class, and return that review
    rev = Review(reviewAuthor, reviewPosition, reviewCompany, reviewRating, reviewDict, days)
    print "Review created for %s..." % rev.name[0]
    sys.stdout.flush()
    return rev


# Create array of Review objects and populate with our reviews
reviewGuide = []
for num in range(reviewNum):
    reviewGuide.append(findMaterials(links[num]))

# Sort our list based on date posted
#reviewGuideSorted = sorted(reviewGuide, key=attrgetter('day'), reverse=True)

# Create document and insert main heading
doc = docx.Document()
doc.add_heading('Trust Radius Weekly Report', 0)

# Func createPage: page (param): an instance of a Review object
def createPage(page):

    # Insert Review Info
    doc.add_heading(page.name, 1)
    doc.add_heading("%s at %s" % (page.position[0], page.company[0]), 3)
    doc.add_heading(page.day.strftime('%B %d, %Y'), 3)
    doc.add_heading(page.rating + ' out of 10 stars', 3)

    # Insert Review Text
    for x, y in page.goodies.items():
        doc.add_heading(x, 4)
        doc.add_paragraph(y)

    # Create new page for next review    
    if reviewNum != 1:
        doc.add_page_break()

# Iterate through all of our reviews to create docx
for review in reviewGuide:
    createPage(review)

# Print success and save docx (if used for plural/singular case)
if reviewNum == 1:
    print 'Successfully created a .docx with %d review. Check out TrustRadiusReport.docx...' % reviewNum
else:
    print 'Successfully created a .docx with %d reviews. Check out TrustRadiusReport.docx...' % reviewNum
    
doc.save('TrustRadiusReport.docx')
